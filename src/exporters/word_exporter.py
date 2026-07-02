"""Word exporter — structured monthly report in .docx format."""
import datetime
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.config import Config

logger = logging.getLogger(__name__)


class WordExporter:
    """Export monthly report to Word (.docx)."""

    def __init__(self, output_dir: str = None):
        self.output_dir = Path(output_dir or Config().export_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export(self, stats: dict, summary: dict, reports: list,
               month_str: str, session=None) -> str:
        """Generate Word document. Returns file path."""
        doc = Document()

        # Title
        title = doc.add_heading(f"医药业务月报 — {month_str}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")

        # ── 1. 统计摘要 ──────────────────────────────────────────
        doc.add_heading("一、统计摘要", level=1)
        table = doc.add_table(rows=8, cols=2, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        summary_data = [
            ("指标", "数值"),
            ("拜访总量", str(stats["total_visits"])),
            ("客户覆盖数（去重）", str(stats["unique_customers"])),
            ("活动数量", str(stats["activity_count"])),
            ("跟进任务总数", str(stats["tasks_total"])),
            ("已完成任务", str(stats["tasks_completed"])),
            ("待处理任务", str(stats["tasks_pending"])),
            ("完成率", f"{stats['completion_rate']}%"),
        ]
        for i, (label, value) in enumerate(summary_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            if i == 0:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

        # ── 2. 科室分布 ──────────────────────────────────────────
        doc.add_heading("二、科室分布", level=1)
        dept_table = doc.add_table(rows=1, cols=2, style="Table Grid")
        dept_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        dept_table.rows[0].cells[0].text = "科室"
        dept_table.rows[0].cells[1].text = "拜访次数"
        for cell in dept_table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for dept, count in stats.get("department_distribution", {}).items():
            row = dept_table.add_row()
            row.cells[0].text = dept
            row.cells[1].text = str(count)

        # ── 3. 高频沟通主题 ─────────────────────────────────────
        doc.add_heading("三、高频沟通主题", level=1)
        for t in stats.get("top_topics", [])[:10]:
            doc.add_paragraph(t, style="List Bullet")

        # ── 4. 重点医学问题 ─────────────────────────────────────
        doc.add_heading("四、重点医学问题", level=1)
        for q in stats.get("medical_questions", [])[:10]:
            doc.add_paragraph(q, style="List Bullet")

        # ── 5. 工作进展（AI 总结） ──────────────────────────────
        doc.add_heading("五、本月工作进展", level=1)
        doc.add_paragraph(summary.get("progress_summary", ""))

        # ── 6. 重点医学问题分析 ─────────────────────────────────
        doc.add_heading("六、重点医学问题分析", level=1)
        doc.add_paragraph(summary.get("key_issues", ""))

        # ── 7. 未完成事项 ────────────────────────────────────────
        doc.add_heading("七、未完成事项", level=1)
        doc.add_paragraph(summary.get("unfinished_items", ""))
        for item in stats.get("pending_items", []):
            doc.add_paragraph(
                f"{item.get('customer', '')}：{item.get('task', '')}"
                f"（截止：{item.get('deadline', '未设')}）",
                style="List Bullet"
            )

        # ── 8. 下月计划 ──────────────────────────────────────────
        doc.add_heading("八、下月工作计划", level=1)
        doc.add_paragraph(summary.get("next_month_plan", ""))

        # ── 9. 日报明细表 ────────────────────────────────────────
        doc.add_heading("九、本月日报明细", level=1)
        if reports:
            detail_table = doc.add_table(rows=1, cols=7, style="Table Grid")
            detail_headers = ["日期", "科室", "客户", "主题", "反馈", "后续任务", "状态"]
            for i, h in enumerate(detail_headers):
                detail_table.rows[0].cells[i].text = h
                for p in detail_table.rows[0].cells[i].paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(8)

            for r in reports[:100]:  # Limit rows
                row = detail_table.add_row()
                cells = [
                    str(r.date) if r.date else "",
                    r.department or "",
                    r.customer_id or "",
                    (r.topic or "")[:60],
                    (r.feedback or "")[:80],
                    (r.follow_up_task or "")[:60],
                    r.follow_up_status or "",
                ]
                for i, val in enumerate(cells):
                    row.cells[i].text = val
                    for p in row.cells[i].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(8)

        # ── 10. 文献附录 ─────────────────────────────────────────
        doc.add_heading("十、文献附录", level=1)
        try:
            from src.db import get_session
            from src.models.literature import LiteratureArticle
            s = session or get_session()
            try:
                articles = s.query(LiteratureArticle).order_by(LiteratureArticle.year.desc()).limit(20).all()
                if articles:
                    for a in articles:
                        p = doc.add_paragraph()
                        p.add_run(f"[{a.pmid or 'N/A'}] ").bold = True
                        p.add_run(f"{a.title or '无标题'} ")
                        p.add_run(f"{a.authors or ''} ").italic = True
                        p.add_run(f"{a.journal or ''}, {a.year or ''}")
                        if a.source_url:
                            p.add_run(f"\n来源：{a.source_url}")
                else:
                    doc.add_paragraph("本月无文献检索记录。")
            finally:
                if not session:
                    s.close()
        except Exception:
            doc.add_paragraph("文献数据暂时不可用。")

        # Disclaimer
        doc.add_paragraph("")
        disclaimer = doc.add_paragraph()
        disclaimer.add_run("免责声明：").bold = True
        disclaimer.add_run("本文档内容仅用于内部信息整理，不构成医学建议。重要结论请务必核验原文。")

        # Save
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"月报_{month_str}_{ts}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)
        logger.info(f"Word exported to {filepath}")
        return str(filepath)
